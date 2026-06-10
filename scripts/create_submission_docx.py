from __future__ import annotations

import os
import sys
from datetime import datetime
from pathlib import Path

ENV_PACKAGES = Path("E:/envment/python-packages")
if ENV_PACKAGES.exists():
    sys.path.insert(0, str(ENV_PACKAGES))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "AdVivid_AI_比赛提交材料整合版.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(97, 108, 122)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_run_font(run, *, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, *, size=11, color=INK, bold=False, name="Microsoft YaHei"):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold, name=name)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, color=INK, size=10.5, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.style = "Heading 1"
    elif level == 2:
        p.style = "Heading 2"
    else:
        p.style = "Heading 3"
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text="", *, bold=False, color=INK, size=11, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=10.8, color=INK)


def next_numbering_id(numbering, attr_name):
    values = []
    for element in numbering.iter():
        value = element.get(qn(f"w:{attr_name}"))
        if value is not None and value.isdigit():
            values.append(int(value))
    return (max(values) + 1) if values else 1


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "abstractNumId")
    num_id = next_numbering_id(numbering, "numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level_type)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=10.8, color=INK)


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.55, 4.95])
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), label, bold=True, color=DARK_BLUE, size=10.5, fill=LIGHT_GRAY)
        set_cell_text(table.cell(idx, 1), value, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color=DARK_BLUE, size=10, fill=LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "AdVivid AI | 电商 AIGC 带货视频生成系统"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_font(header, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "比赛提交材料整合版"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(footer, size=9, color=MUTED)
    return doc


def add_cover(doc):
    add_para(doc, "比赛提交材料整合版", bold=True, color=BLUE, size=12, after=12)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.0
    r = title.add_run("AdVivid AI：电商带货视频智能创作系统")
    set_run_font(r, size=25, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    r2 = subtitle.add_run("面向商家的 AIGC 短视频创作工作台")
    set_run_font(r2, size=14, color=MUTED)

    add_key_value_table(
        doc,
        [
            ("参赛课题", "电商场景 AIGC 带货视频生成系统"),
            ("文档用途", "提交表信息整合、答辩说明、演示视频脚本与评审材料索引"),
            ("当前状态", "P0 已完成，P1 主要能力可演示，P2 已预留扩展路径"),
            ("待补充项", "团队成员与分工、在线 Demo 链接、演示视频链接、源代码仓库链接"),
            ("生成日期", datetime.now().strftime("%Y-%m-%d")),
        ],
    )

    add_callout(
        doc,
        "一句话业务价值",
        "帮助商家从商品素材和卖点出发，自动生成可编辑、可导出、可复盘的短视频带货内容，降低内容生产成本并提升投放迭代效率。",
        fill="EEF6F6",
    )

    add_para(
        doc,
        "本文档根据赛题交付要求整理，覆盖项目基础信息、核心功能、端到端流程、系统架构、AI Agent、工程实现、部署说明、演示脚本和最终提交清单。",
        color=MUTED,
        size=10.5,
    )
    doc.add_page_break()


def add_submission_info(doc):
    add_heading(doc, "1. 基础信息与提交项", 1)
    add_key_value_table(
        doc,
        [
            ("项目名称", "AdVivid AI：电商带货视频智能创作系统"),
            ("参赛课题", "电商场景 AIGC 带货视频生成系统"),
            ("一句话价值", "从商品素材和卖点自动生成可编辑、可导出、可复盘的短视频带货内容。"),
            ("在线 Demo", "待填写：部署后的 Web 地址"),
            ("演示视频", "待填写：3-5 分钟演示视频链接"),
            ("源代码仓库", "待填写：Git 仓库链接"),
            ("团队成员与分工", "待填写：成员 A/B/C 或实际成员姓名与职责"),
        ],
    )
    add_heading(doc, "需要最终手动补齐的字段", 2)
    add_bullets(
        doc,
        [
            "团队成员与分工：建议按前端、后端、AI Agent/视频生成、文档与演示四类写清楚。",
            "在线 Demo 链接：建议提供免登录或体验账号，并确保评委能直接访问。",
            "演示视频链接：建议 3-5 分钟，覆盖素材上传、Agent trace、分镜编辑、一键成片和数据看板。",
            "源代码仓库链接：提交前确认 `.env`、API Key、运行时产物和本地视频文件未进入公开仓库。",
        ]
    )


def add_completion(doc):
    add_heading(doc, "2. 项目定位与完成状态", 1)
    add_callout(
        doc,
        "项目定位",
        "不是简单的视频生成工具，而是一个面向商家的电商短视频智能创作工作台。系统强调素材管理、剧本生成、分镜可控、长任务稳定、视频导出和数据复盘的一体化流程。",
    )
    add_matrix_table(
        doc,
        ["层级", "要求", "当前完成情况"],
        [
            ("P0", "素材上传、剧本生成、基础分镜、一键成片、任务进度、预览导出", "已完成，可端到端演示。"),
            ("P1", "素材标签/Embedding、Agent、分镜编辑、TTS/字幕/BGM、失败重试、Trace、Mock 看板", "主要能力已完成，可作为作品亮点展示。"),
            ("P2", "A/B、多因子归因、CI/CD、可观测性、合规审核流", "已预留设计思路，其中 CI 已实现基础流水线。"),
        ],
        [0.8, 2.8, 2.9],
    )


def add_features(doc):
    add_heading(doc, "3. 核心功能", 1)
    add_matrix_table(
        doc,
        ["功能", "用户效果", "实现要点"],
        [
            (
                "商品项目与素材库",
                "商家可录入商品卖点、人群、场景和创作补充要求，上传商品图/视频/参考素材。",
                "React 工作台 + Fastify 上传 API；素材记录保存到 JSON Store 或 Prisma/PostgreSQL。",
            ),
            (
                "素材结构化与检索",
                "素材可被标签化、摘要化、切片化，并在分镜中被推荐使用。",
                "图片/视频分析、FFmpeg 缩略图、mock embedding、关键词+标签+向量混合检索。",
            ),
            (
                "LangGraph 创作 Agent",
                "AI 不只生成文案，而是参与商品理解、素材召回、策略选择、剧本与分镜规划。",
                "Python FastAPI + LangChain + LangGraph + Pydantic；每个节点写入 generation_traces。",
            ),
            (
                "分镜级编辑",
                "用户可修改单个分镜台词、字幕、画面、镜头、时长、素材绑定和顺序。",
                "scenes 独立持久化；支持单镜重生成和单镜预览。",
            ),
            (
                "一键成片",
                "输出约 15-20 秒、30 秒以内的短视频，支持预览和下载。",
                "Seedance 分段生成 + FFmpeg 拼接；失败时 FFmpeg 本地兜底，含字幕、BGM、mock TTS。",
            ),
            (
                "任务与看板",
                "长任务有状态、进度、失败原因和重试；看板展示数据回流思路。",
                "local runner 或 BullMQ/Redis；ECharts mock 数据展示播放、CTR、转化和创作因子。",
            ),
        ],
        [1.35, 2.25, 2.9],
    )


def add_flow(doc):
    add_heading(doc, "4. 端到端使用流程", 1)
    add_numbers(
        doc,
        [
            "商家打开 AdVivid AI 工作台，新建一个商品视频项目。",
            "录入商品标题、核心卖点、目标人群、使用场景、创作策略、目标时长和 Prompt 微调要求。",
            "进入素材库上传商品主图、商品视频或参考素材，系统自动分析素材摘要、标签、embedding 和视频切片。",
            "点击生成剧本后，Node.js API 创建异步任务，Python LangGraph Agent 依次完成商品理解、素材召回、策略选择、剧本生成、分镜规划和质量检查。",
            "任务页展示状态、进度和生成 trace，商家可以看到 AI 每一步做了什么。",
            "剧本生成完成后，商家回到创作台查看完整带货剧本和分镜列表，并对单个分镜做局部编辑。",
            "商家可对某个分镜单独重生成或单镜预览，确认局部效果后再发起整片渲染。",
            "系统生成视频，支持 9:16/16:9、720p/1080p 导出，并用 Mock 看板展示后续投放优化思路。",
        ]
    )


def add_architecture(doc):
    add_heading(doc, "5. 系统架构与技术栈", 1)
    add_para(doc, "总体架构采用前后端分离 + 长任务队列 + AI Agent 编排 + 视频渲染流水线。", color=MUTED)
    add_matrix_table(
        doc,
        ["层级", "技术选择", "职责"],
        [
            ("前端", "React + Vite + TypeScript + TailwindCSS + ECharts", "商家工作台、素材管理、分镜编辑、任务进度、视频预览和数据看板。"),
            ("后端 API", "Node.js + TypeScript + Fastify", "HTTP API、文件上传、数据库读写、任务创建、模型调度和视频导出。"),
            ("AI Agent", "Python + FastAPI + LangChain + LangGraph", "商品理解、素材召回、策略选择、剧本生成、分镜规划、质检和渲染计划。"),
            ("任务队列", "local runner 或 BullMQ + Redis", "处理剧本生成、素材分析和视频渲染等长任务。"),
            ("数据存储", "JSON Store 或 Prisma + PostgreSQL/pgvector", "本地零配置演示与生产级扩展两种路径。"),
            ("视频处理", "Seedance + FFmpeg", "真实视频生成、分段拼接、字幕、BGM、mock TTS 与兜底出片。"),
            ("部署", "Docker Compose + Nginx", "Web、API、Worker、Python Agent、PostgreSQL、Redis 一体部署。"),
        ],
        [1.1, 2.35, 3.05],
    )
    add_heading(doc, "架构流程", 2)
    add_para(
        doc,
        "React Web -> Fastify API -> JSON/Prisma Store -> Local Runner/BullMQ -> Python LangGraph Agent -> Ark/Seedance/FFmpeg -> 视频预览与导出",
        bold=True,
        color=DARK_BLUE,
        size=10.5,
    )


def add_agent(doc):
    add_heading(doc, "6. AI Agent 与模型能力", 1)
    add_matrix_table(
        doc,
        ["Agent 节点", "作用", "输出"],
        [
            ("ProductAnalyzer", "理解商品标题、卖点、人群、场景和创作补充要求。", "ProductProfile、关键词、核心利益点。"),
            ("MaterialRetriever", "从素材库和切片中召回候选素材。", "候选素材、切片 ID、匹配理由。"),
            ("StrategySelector", "选择痛点开场、场景种草、测评对比等创作策略。", "creativeStrategy、hook 类型、创作因子。"),
            ("ScriptWriter", "生成完整带货剧本摘要。", "标题、叙事结构、hook、约束清单。"),
            ("ScenePlanner", "拆成 5-6 个连续分镜。", "画面、镜头、台词、字幕、时长、素材建议。"),
            ("ReviewAgent", "检查时长、分镜数量、字幕和商品一致性。", "质检结果、问题与建议。"),
            ("RenderPlanner", "将分镜转为可执行渲染计划。", "renderPlan、音频和字幕参数。"),
        ],
        [1.45, 2.55, 2.5],
    )
    add_callout(
        doc,
        "AI 亮点表达",
        "大模型在本项目中不是一个黑盒文案生成器，而是作为可追踪的创作 Agent 参与商品理解、素材匹配、创意策略、分镜规划和质量检查。每一步都有 trace，既能展示 AI 过程，也便于失败定位和重试。",
    )
    add_heading(doc, "模型与兜底策略", 2)
    add_bullets(
        doc,
        [
            "文本生成：优先调用火山方舟 Doubao 文本模型，失败或超时时走 Python/TypeScript 连续剧情模板兜底。",
            "视频生成：优先调用 Seedance，支持 5 秒片段分段生成并拼接为约 15-20 秒视频。",
            "本地渲染：Seedance 失败时自动回退 FFmpeg，仍能生成可播放、可下载的 mp4。",
            "安全边界：API Key 只存在 `.env`，模型状态接口只返回是否配置，不返回 Key 或 endpoint 详情。",
        ]
    )


def add_requirements(doc):
    add_heading(doc, "7. 赛题要求对照", 1)
    add_matrix_table(
        doc,
        ["赛题要求", "当前实现", "完成度"],
        [
            ("React / Node.js / TypeScript", "React 前端工作台、Fastify API、全栈 TypeScript 类型约束。", "已完成"),
            ("火山 OpenAPI 与开源框架", "Ark 文本、Seedance 视频、Python LangGraph、FFmpeg。", "已完成"),
            ("素材导入与管理", "图片/视频上传、预览、删除、重新分析、切片、标签、摘要。", "已完成"),
            ("Prompt 调整", "创作台支持“创作补充要求 / Prompt 微调”，进入 Agent Prompt。", "已完成"),
            ("分镜干预", "修改台词、字幕、画面、镜头、时长、素材绑定、排序、单镜重生成。", "已完成"),
            ("高质量视频生成", "Seedance 接入，FFmpeg 兜底，字幕/BGM/mock TTS，15-20 秒成片。", "已完成"),
            ("进度、重试、兜底", "render_jobs 状态、进度、错误、retry、trace；local/BullMQ 双模式。", "已完成"),
            ("数据反馈", "Mock 看板展示 Hook、风格因子、播放、CTR、转化。", "可演示"),
            ("CI/CD 与工程化", "ESLint、Prettier、StyleLint、Husky、GitHub Actions CI。", "已补齐"),
            ("合规与安全", "素材授权提示、Key 不暴露、`.env` 不提交、模型输出 schema 校验。", "已覆盖"),
        ],
        [1.75, 3.75, 1.0],
    )


def add_data_api(doc):
    add_heading(doc, "8. 数据库与 API 摘要", 1)
    add_heading(doc, "核心数据实体", 2)
    add_bullets(
        doc,
        [
            "products：商品基础信息、卖点、人群、场景、风格、创作补充要求。",
            "materials / material_slices：上传素材、视频切片、摘要、标签、embedding。",
            "projects：一次完整视频创作项目。",
            "scripts / scenes：结构化剧本和分镜脚本。",
            "render_jobs：剧本生成、素材分析、视频渲染和单镜预览任务。",
            "generated_videos：生成视频版本、比例、清晰度和导出路径。",
            "generation_traces：Agent 与 Worker 每一步执行记录。",
            "analytics_mock_events：Mock 播放、点击、转化数据。",
        ]
    )
    add_heading(doc, "核心 API 分组", 2)
    add_matrix_table(
        doc,
        ["分组", "接口示例", "说明"],
        [
            ("素材", "POST /api/materials, GET /api/materials/search", "上传、分析、搜索、删除素材。"),
            ("项目/商品", "POST /api/projects, PATCH /api/products/:id", "创建项目和保存商品信息。"),
            ("剧本/分镜", "POST /api/scripts/generate, PATCH /api/scenes/:id", "生成剧本、编辑分镜、排序、单镜重生成。"),
            ("视频", "POST /api/videos/render, GET /api/videos/:id/export", "创建渲染任务、预览和导出成片。"),
            ("任务", "GET /api/jobs/:id, POST /api/jobs/:id/retry", "查询进度、失败重试、查看 trace。"),
            ("看板", "GET /api/analytics/mock", "展示 Mock 数据回流与创作因子效果。"),
        ],
        [1.0, 2.7, 2.8],
    )


def add_engineering(doc):
    add_heading(doc, "9. 工程难点与解决方案", 1)
    add_matrix_table(
        doc,
        ["难点", "风险", "解决方案"],
        [
            ("长任务耗时", "HTTP 阻塞、用户不知道进度。", "任务表 + local/BullMQ 队列 + 前端轮询 + 进度和当前步骤。"),
            ("外部模型不稳定", "Key、网络、配额会影响演示。", "Ark/Seedance 均有本地兜底；Seedance 失败回退 FFmpeg。"),
            ("视频剧情割裂", "片段生硬拼接，带货逻辑弱。", "Agent 与兜底模板都要求开场、商品介入、卖点证明、使用结果、轻 CTA。"),
            ("黑盒不可控", "商家无法改局部内容。", "剧本和分镜结构化保存，支持单镜编辑、重生成、预览和重新渲染。"),
            ("素材复用", "上传素材难以自动匹配。", "素材摘要、标签、切片、mock embedding 与本地混合检索推荐。"),
            ("安全合规", "Key 泄露或素材版权问题。", "Key 只放 `.env`，前端不暴露；素材页提示使用自有或授权素材。"),
        ],
        [1.4, 1.95, 3.15],
    )


def add_deployment(doc):
    add_heading(doc, "10. 部署、访问与安全", 1)
    add_heading(doc, "本地运行", 2)
    add_para(doc, "npm install\nnpm run setup:python\nnpm run dev", size=10.5, color=DARK_BLUE, bold=True)
    add_heading(doc, "访问地址", 2)
    add_bullets(
        doc,
        [
            "Web：http://localhost:5173",
            "API：http://localhost:4000/api/health",
            "Python Agent：http://localhost:8002/health",
        ]
    )
    add_heading(doc, "生产 Demo", 2)
    add_para(
        doc,
        "docker compose --env-file .env -f infra/docker/docker-compose.prod.yml up --build -d",
        size=10.5,
        color=DARK_BLUE,
        bold=True,
    )
    add_heading(doc, "安全说明", 2)
    add_bullets(
        doc,
        [
            "真实 API Key 只写入 `.env`，不写入 `.env.example`、README、截图或演示视频。",
            "前端不会拿到模型 Key，所有模型调用均由服务端完成。",
            "新增数据库、缓存或运行时依赖在 Windows 本地统一放到 `E:/envment`。",
            "上传文件限制类型和大小，模型输出通过 Pydantic/Zod schema 校验。",
        ]
    )


def add_demo(doc):
    add_heading(doc, "11. 演示视频脚本与截图清单", 1)
    add_heading(doc, "推荐 3-5 分钟演示脚本", 2)
    add_numbers(
        doc,
        [
            "打开首页，说明这是商家短视频智能创作工作台。",
            "新建商品项目，填写商品标题、卖点、人群、场景、策略和创作补充要求。",
            "上传商品图或视频，展示素材分析、标签、摘要、切片和预览。",
            "点击生成剧本，进入任务页展示 LangGraph trace。",
            "回到创作台展示结构化剧本和 5-6 个连续分镜。",
            "编辑一个分镜的字幕或时长，并绑定推荐素材。",
            "点击单镜预览，说明支持局部检查与局部重生成。",
            "选择导出比例和清晰度，点击一键成片。",
            "预览生成视频，展示字幕、BGM、画面和下载。",
            "打开数据看板，说明 Mock 数据如何反哺创作因子优化。",
        ]
    )
    add_heading(doc, "建议截图", 2)
    add_bullets(
        doc,
        [
            "创作台商品信息与 Prompt 微调字段。",
            "素材库上传、预览、标签和切片。",
            "任务进度与 LangGraph trace。",
            "分镜编辑器与单镜预览。",
            "视频预览与导出参数。",
            "Mock 数据看板。",
        ]
    )


def add_final_checklist(doc):
    add_heading(doc, "12. 最终提交清单", 1)
    add_matrix_table(
        doc,
        ["提交材料", "当前准备情况", "提交前动作"],
        [
            ("项目名称", "已确定：AdVivid AI：电商带货视频智能创作系统", "直接填写"),
            ("参赛课题", "已确定：电商场景 AIGC 带货视频生成系统", "直接填写"),
            ("团队成员与分工", "文档中已预留", "补真实姓名和职责"),
            ("一句话业务价值", "已整理", "直接填写"),
            ("在线 Demo 链接", "待部署后填写", "确认可访问、可演示"),
            ("演示视频链接", "待录制后填写", "上传并确认权限"),
            ("源代码仓库链接", "待填写", "提交前检查不泄露 Key"),
            ("README / 运行说明", "已存在", "确认命令可复现"),
            ("架构图/API/数据库/部署说明", "docs 中已整理", "提交时作为附件或仓库文档"),
            ("生成样例与截图", "待最终录制/截图", "准备 1-2 条视频样例"),
        ],
        [1.5, 2.6, 2.4],
    )
    add_callout(
        doc,
        "答辩核心表达",
        "本项目围绕电商商家的真实短视频生产流程，构建了素材管理、剧本生成、分镜创作、视频渲染、数据复盘的一体化系统。系统使用 React、Node.js 和 TypeScript 实现全栈工程架构，并引入 Python LangGraph 编排电商创作 Agent，使大模型参与商品理解、素材召回、创意策略选择、分镜规划和质量检查，最终形成可编辑、可追踪、可导出的 AIGC 带货视频生产链路。",
        fill="EEF6F6",
    )


def main():
    doc = setup_document()
    add_cover(doc)
    add_submission_info(doc)
    add_completion(doc)
    add_features(doc)
    add_flow(doc)
    add_architecture(doc)
    add_agent(doc)
    add_requirements(doc)
    add_data_api(doc)
    add_engineering(doc)
    add_deployment(doc)
    add_demo(doc)
    add_final_checklist(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
